import streamlit as st
import json
import openai
from openai import OpenAI
import io
import re
import os
import time
import docx
from pypdf import PdfReader
from prompts import PROMPT_GPT_TABLA_PLANIFICACION

# Importamos las funciones que necesitamos de nuestros otros módulos
from drive_utils import (
    find_or_create_folder, get_files_in_project, delete_file_from_drive,
    upload_file_to_drive, find_file_by_name, download_file_from_drive,
    sync_guiones_folders_with_index, list_ect_folders, ROOT_FOLDER_NAME
)
from utils import (
    mostrar_indice_desplegable, limpiar_respuesta_json, agregar_markdown_a_word,
    wrap_html_fragment, html_a_imagen, limpiar_respuesta_final,
    corregir_numeracion_markdown, generar_indice_word
)


from utils import (
    mostrar_indice_desplegable, limpiar_respuesta_json, agregar_markdown_a_word,
    wrap_html_fragment, html_a_imagen, limpiar_respuesta_final,
    corregir_numeracion_markdown, generar_indice_word,
    natural_sort_key  # <--- ¡AÑADE ESTO!
)
# =============================================================================
#           PÁGINA DE BIENVENIDA / INICIO DE SESIÓN
# =============================================================================

def landing_page():
    """Pantalla de bienvenida que incluye el inicio de sesión con Google."""
    # Importación local para evitar dependencias circulares
    from auth import get_google_flow

    col1, col_center, col3 = st.columns([1, 2, 1])
    with col_center:
        st.write("")
        st.markdown(f'<div style="text-align: center;"><img src="https://raw.githubusercontent.com/soporte2-tech/appfront/main/imagen.png" width="150"></div>', unsafe_allow_html=True)
        st.write("")
        st.markdown("<div style='text-align: center;'><h1>Asistente Inteligente para Memorias Técnicas</h1></div>", unsafe_allow_html=True)
        st.markdown("<div style='text-align: center;'><h3>Optimiza y acelera la creación de tus propuestas de licitación</h3></div>", unsafe_allow_html=True)
        st.markdown("---")
        st.info("Para empezar, necesitas dar permiso a la aplicación para que gestione los proyectos en tu Google Drive.")
        
        flow = get_google_flow()
        auth_url, _ = flow.authorization_url(prompt='consent')
        
        st.link_button("🔗 Conectar con Google Drive", auth_url, use_container_width=True, type="primary")

def project_selection_page():
    from app import go_to_landing, go_to_phase1
    """Página para seleccionar o crear un proyecto en Google Drive."""
    st.markdown("<h3>Selección de Proyecto</h3>", unsafe_allow_html=True)
    st.markdown("Elige un proyecto existente de tu Google Drive o crea uno nuevo para empezar.")
    st.markdown("---")
    
    service = st.session_state.drive_service
    if not service:
        st.error("No se pudo conectar con Google Drive. Por favor, intenta volver a la página de inicio y reconectar.")
        if st.button("← Volver a Inicio"):
            for key in ['credentials', 'drive_service']:
                if key in st.session_state: del st.session_state[key]
            go_to_landing(); st.rerun()
        return

    with st.spinner("Accediendo a tu Google Drive..."):
        root_folder_id = find_or_create_folder(service, ROOT_FOLDER_NAME)
        projects = list_project_folders(service, root_folder_id)
    
    with st.container(border=True):
        st.subheader("1. Elige un proyecto existente")
        if not projects:
            st.info("Aún no tienes proyectos. Crea uno nuevo en el paso 2.")
        else:
            project_names = ["-- Selecciona un proyecto --"] + list(projects.keys())
            selected_name = st.selectbox(
                "Selecciona tu proyecto:",
                project_names, 
                key="project_selector"  # <--- AÑADE ESTA LÍNEA
            )
            
            if st.button("Cargar Proyecto Seleccionado", type="primary"):
                if selected_name != "-- Selecciona un proyecto --":
                    st.session_state.selected_project = {"name": selected_name, "id": projects[selected_name]}
                    st.toast(f"Proyecto '{selected_name}' cargado."); go_to_phase1(); st.rerun()
                else:
                    st.warning("Por favor, selecciona un proyecto de la lista.")

    with st.container(border=True):
        st.subheader("2. O crea un nuevo proyecto")
        new_project_name = st.text_input("Nombre del nuevo proyecto (ej: Licitación Metro Madrid 2024)", key="new_project_name_input")
        if st.button("Crear y Empezar Nuevo Proyecto"):
            if not new_project_name.strip():
                st.warning("Por favor, introduce un nombre para el proyecto.")
            elif new_project_name in projects:
                st.error("Ya existe un proyecto con ese nombre. Por favor, elige otro.")
            else:
                with st.spinner(f"Creando carpeta '{new_project_name}' en tu Drive..."):
                    new_project_id = find_or_create_folder(service, new_project_name, parent_id=root_folder_id)
                    st.session_state.selected_project = {"name": new_project_name, "id": new_project_id}
                    st.success(f"¡Proyecto '{new_project_name}' creado! Ya puedes cargar los documentos.")
                    go_to_phase1(); st.rerun()

# =============================================================================
#           FASE 1: ANÁLISIS Y ESTRUCTURA
# =============================================================================

def phase_1_page(model):
    from app import handle_full_regeneration, go_to_project_selection, go_to_phase1_results, back_to_project_selection_and_cleanup
    """Página de Fase 1 para gestionar archivos y generar el índice."""

    if not st.session_state.get('selected_project'):
        st.warning("No se ha seleccionado ningún proyecto. Volviendo a la selección.")
        go_to_project_selection(); st.rerun()

    project_name = st.session_state.selected_project['name']
    project_folder_id = st.session_state.selected_project['id']
    service = st.session_state.drive_service

    st.markdown(f"<h3>FASE 1: Análisis y Estructura</h3>", unsafe_allow_html=True)
    st.info(f"Estás trabajando en el proyecto: **{project_name}**")

    pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
    document_files = get_files_in_project(service, pliegos_folder_id)
    
    if document_files:
        st.success("Hemos encontrado estos archivos en la carpeta 'Pliegos' de tu proyecto:")
        with st.container(border=True):
            for file in document_files:
                cols = st.columns([4, 1])
                cols[0].write(f"📄 **{file['name']}**")
                if cols[1].button("Eliminar", key=f"del_{file['id']}", type="secondary"):
                    with st.spinner(f"Eliminando '{file['name']}'..."):
                        if delete_file_from_drive(service, file['id']):
                            st.toast(f"Archivo '{file['name']}' eliminado."); st.rerun()
    else:
        st.info("La carpeta 'Pliegos' de este proyecto está vacía. Sube los archivos base.")

    with st.expander("Añadir o reemplazar documentación en la carpeta 'Pliegos'", expanded=not document_files):
        with st.container(border=True):
            st.subheader("Subir nuevos documentos")
            new_files_uploader = st.file_uploader("Arrastra aquí los nuevos Pliegos o Plantilla", type=['docx', 'pdf'], accept_multiple_files=True, key="new_files_uploader")
            if st.button("Guardar nuevos archivos en Drive"):
                if new_files_uploader:
                    with st.spinner("Subiendo archivos a la carpeta 'Pliegos'..."):
                        for file_obj in new_files_uploader:
                            upload_file_to_drive(service, file_obj, pliegos_folder_id)
                        st.rerun()
                else:
                    st.warning("Por favor, selecciona al menos un archivo para subir.")

    st.markdown("---"); st.header("Análisis y Generación de Índice")
    
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
    saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Cargar último índice generado", use_container_width=True, disabled=not saved_index_id):
            with st.spinner("Cargando índice desde Drive..."):
                index_content_bytes = download_file_from_drive(service, saved_index_id)
                index_data = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.session_state.generated_structure = index_data
                st.session_state.uploaded_pliegos = document_files
                go_to_phase1_results(); st.rerun()

    with col2:
        if st.button("Analizar Archivos y Generar Nuevo Índice", type="primary", use_container_width=True, disabled=not document_files):
            if handle_full_regeneration(model):
                go_to_phase1_results(); st.rerun()

    st.write(""); st.markdown("---")
    st.button("← Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup, use_container_width=True, key="back_to_projects")

def phase_1_results_page(model):
    from app import handle_full_regeneration, go_to_phase1, go_to_phase2
    from prompts import PROMPT_REGENERACION
    """Página para revisar, regenerar y aceptar el índice."""
    # Importación local para evitar dependencias circulares
    st.markdown("<h3>FASE 1: Revisión de Resultados</h3>", unsafe_allow_html=True)
    st.markdown("Revisa el índice. Puedes hacer ajustes con feedback, regenerarlo todo desde cero, o aceptarlo para continuar.")
    st.markdown("---")
    st.button("← Volver a la gestión de archivos", on_click=go_to_phase1)

    if 'generated_structure' not in st.session_state or not st.session_state.generated_structure:
        st.warning("No se ha generado ninguna estructura.")
        return

    def handle_regeneration_with_feedback():
        feedback_text = st.session_state.feedback_area
        if not feedback_text:
            st.warning("Por favor, escribe tus indicaciones en el área de texto.")
            return

        with st.spinner("🧠 Incorporando tu feedback y regenerando la estructura..."):
            try:
                contenido_ia_regeneracion = [PROMPT_REGENERACION]
                contenido_ia_regeneracion.append("--- INSTRUCCIONES DEL USUARIO ---\n" + feedback_text)
                contenido_ia_regeneracion.append("--- ESTRUCTURA JSON ANTERIOR A CORREGIR ---\n" + json.dumps(st.session_state.generated_structure, indent=2))
                
                if st.session_state.get('uploaded_pliegos'):
                    service = st.session_state.drive_service
                    for file_info in st.session_state.uploaded_pliegos:
                        file_content_bytes = download_file_from_drive(service, file_info['id'])
                        contenido_ia_regeneracion.append({"mime_type": file_info['mimeType'], "data": file_content_bytes.getvalue()})

                generation_config = genai.GenerationConfig(response_mime_type="application/json")
                response_regeneracion = model.generate_content(contenido_ia_regeneracion, generation_config=generation_config)
                json_limpio_str_regenerado = limpiar_respuesta_json(response_regeneracion.text)
                
                if json_limpio_str_regenerado:
                    st.session_state.generated_structure = json.loads(json_limpio_str_regenerado)
                    st.toast("¡Estructura regenerada con feedback!")
                    st.session_state.feedback_area = "" # Limpiamos el área de texto
                else:
                    st.error("La IA no devolvió una estructura válida tras la regeneración.")
            except Exception as e:
                st.error(f"Ocurrió un error durante la regeneración: {e}")

    with st.container(border=True):
        mostrar_indice_desplegable(st.session_state.generated_structure.get('estructura_memoria'))
        st.markdown("---")
        st.subheader("Validación y Siguiente Paso")
        
        st.text_area("Si necesitas cambios, indícalos aquí:", key="feedback_area", placeholder="Ej: 'Une los apartados 1.1 y 1.2 en uno solo.'")
        
        col1, col2 = st.columns(2)
        with col1:
            st.button(
                "Regenerar con Feedback", 
                on_click=handle_regeneration_with_feedback, 
                use_container_width=True, 
                disabled=not st.session_state.get("feedback_area")
            )
        with col2:
            st.button(
                "🔁 Regenerar Índice Entero", 
                on_click=handle_full_regeneration, 
                args=(model,), 
                use_container_width=True, 
                help="Descarta este índice y genera uno nuevo desde cero analizando los pliegos otra vez."
            )

        if st.button("Aceptar Índice y Pasar a Fase 2 →", type="primary", use_container_width=True):
            with st.spinner("Sincronizando carpetas y guardando índice final en Drive..."):
                try:
                    service = st.session_state.drive_service
                    project_folder_id = st.session_state.selected_project['id']
                    
                    deleted_count = sync_guiones_folders_with_index(service, project_folder_id, st.session_state.generated_structure)
                    if deleted_count > 0:
                        st.success(f"Limpieza completada: {deleted_count} carpetas de guiones obsoletas eliminadas.")

                    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
                    json_bytes = json.dumps(st.session_state.generated_structure, indent=2).encode('utf-8')
                    mock_file_obj = io.BytesIO(json_bytes)
                    mock_file_obj.name = "ultimo_indice.json"
                    mock_file_obj.type = "application/json"
                    
                    saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
                    if saved_index_id:
                        delete_file_from_drive(service, saved_index_id)
                    upload_file_to_drive(service, mock_file_obj, docs_app_folder_id)
                    st.toast("Índice final guardado en tu proyecto de Drive.")
                    
                    go_to_phase2()
                    st.rerun()

                except Exception as e:
                    st.error(f"Ocurrió un error durante la sincronización o guardado: {e}")

def phase_2_page(model):
    """
    Centro de mando para la generación de guiones.
    MODIFICADO: Permite cambiar entre el motor de Gemini y GPT con una variable.
    """
    from app import go_to_phase1, go_to_phase1_results, go_to_phase4, go_to_phase3
    from prompts import PROMPT_PREGUNTAS_TECNICAS_INDIVIDUAL

    # =================== [ CONTROL PARA PRUEBAS ] ===================
    USE_GPT_MODEL = True
    # ================================================================

    st.markdown("<h3>FASE 2: Centro de Mando de Guiones</h3>", unsafe_allow_html=True)
    st.markdown("Gestiona tus guiones de forma individual o selecciónalos para generarlos en lote.")
    st.markdown("---")

    # --- SETUP INICIAL Y CARGA DE ÍNDICE (sin cambios) ---
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    if 'generated_structure' not in st.session_state:
        st.info("Sincronizando índice desde Google Drive...")
        try:
            docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
            saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
            if saved_index_id:
                index_content_bytes = download_file_from_drive(service, saved_index_id)
                st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.rerun()
            else:
                st.warning("No se ha encontrado un índice guardado. Por favor, vuelve a la Fase 1 para generar uno.")
                if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
                return
        except Exception as e:
            st.error(f"Error al cargar el índice desde Drive: {e}")
            return

    # --- CONSTRUCCIÓN DE LISTA ROBUSTA (sin cambios) ---
    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict) and 'subapartado' in item}
    if not estructura: st.error("La estructura JSON no contiene la clave 'estructura_memoria'."); return
    subapartados_a_mostrar = []
    for seccion in estructura:
        apartado_principal = seccion.get('apartado', 'Sin Título')
        for subapartado_titulo in seccion.get('subapartados', []):
            matiz_existente = matices_dict.get(subapartado_titulo)
            if matiz_existente: subapartados_a_mostrar.append(matiz_existente)
            else: subapartados_a_mostrar.append({"apartado": apartado_principal, "subapartado": subapartado_titulo, "indicaciones": "No se encontraron indicaciones detalladas."})
    if not subapartados_a_mostrar: st.warning("El índice no contiene subapartados."); return

    # --- FUNCIONES DE ACCIÓN INTERNAS (AMBAS DISPONIBLES) ---

    def ejecutar_generacion_con_gemini(model, titulo, indicaciones_completas, show_toast=True):
        """
        Genera el borrador de un subapartado usando el modelo de Gemini.
        """
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", titulo)
        nombre_archivo = nombre_limpio + ".docx"
        try:
            service = st.session_state.drive_service
            project_folder_id = st.session_state.selected_project['id']
            
            guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            subapartado_guion_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_folder_id)
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            pliegos_en_drive = get_files_in_project(service, pliegos_folder_id)
            
            contenido_ia = [PROMPT_PREGUNTAS_TECNICAS_INDIVIDUAL]
            contenido_ia.append("--- INDICACIONES PARA ESTE APARTADO ---\n" + json.dumps(indicaciones_completas, indent=2, ensure_ascii=False))
            
            for file_info in pliegos_en_drive:
                file_content_bytes = download_file_from_drive(service, file_info['id'])
                contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_content_bytes.getvalue()})
                
            doc_extra_key = f"upload_{titulo}"
            if doc_extra_key in st.session_state and st.session_state[doc_extra_key]:
                for uploaded_file in st.session_state[doc_extra_key]:
                    contenido_ia.append("--- DOCUMENTACIÓN DE APOYO ADICIONAL ---\n")
                    contenido_ia.append({"mime_type": uploaded_file.type, "data": uploaded_file.getvalue()})
                    upload_file_to_drive(service, uploaded_file, subapartado_guion_folder_id)
                    
            response = model.generate_content(contenido_ia)
            
            documento = docx.Document()
            agregar_markdown_a_word(documento, response.text)
            doc_io = io.BytesIO()
            documento.save(doc_io)
            word_file_obj = io.BytesIO(doc_io.getvalue())
            word_file_obj.name = nombre_archivo
            word_file_obj.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            upload_file_to_drive(service, word_file_obj, subapartado_guion_folder_id)
            
            if show_toast: 
                st.toast(f"Borrador (Gemini) para '{titulo}' generado y guardado.")
            return True
        
        except Exception as e: 
            st.error(f"Error al generar con Gemini para '{titulo}': {e}")
            return False
        
    def ejecutar_generacion_con_gpt(titulo, indicaciones_completas, show_toast=True):
        try:
            client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        except Exception:
            st.error("Error: 'OPENAI_API_KEY' no encontrada en secrets.toml.")
            return False

        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", titulo)
        nombre_archivo = nombre_limpio + ".docx"
        try:
            contexto_para_gpt = "--- INDICACIONES CLAVE PARA EL SUBAPARTADO ---\n"
            contexto_para_gpt += json.dumps(indicaciones_completas, indent=2, ensure_ascii=False)
            contexto_para_gpt += "\n\n"
            
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            pliegos_en_drive = get_files_in_project(service, pliegos_folder_id)
            contexto_para_gpt += "--- CONTENIDO COMPLETO DE LOS DOCUMENTOS DE REFERENCIA (PLIEGOS) ---\n"
            for file_info in pliegos_en_drive:
                file_content_bytes = download_file_from_drive(service, file_info['id'])
                texto_extraido = ""
                try:
                    if file_info['name'].endswith('.pdf'):
                        reader = PdfReader(io.BytesIO(file_content_bytes.getvalue()))
                        texto_extraido = "\n".join(page.extract_text() for page in reader.pages)
                    elif file_info['name'].endswith('.docx'):
                        doc = docx.Document(io.BytesIO(file_content_bytes.getvalue()))
                        texto_extraido = "\n".join(para.text for para in doc.paragraphs)
                except Exception as e:
                    st.warning(f"No se pudo procesar el archivo '{file_info['name']}': {e}")
                contexto_para_gpt += f"**Inicio del documento: {file_info['name']}**\n{texto_extraido}\n**Fin del documento: {file_info['name']}**\n\n"

            response = client.chat.completions.create(
                model="gpt-4o-mini", # Usando el modelo más reciente y económico
                messages=[
                    {"role": "system", "content": PROMPT_GPT_TABLA_PLANIFICACION},
                    {"role": "user", "content": contexto_para_gpt}
                ],
                temperature=0.2,
            )
            guion_generado = response.choices[0].message.content

            guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            subapartado_guion_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_folder_id)
            documento = docx.Document()
            agregar_markdown_a_word(documento, guion_generado)
            doc_io = io.BytesIO()
            documento.save(doc_io)
            word_file_obj = io.BytesIO(doc_io.getvalue())
            word_file_obj.name = nombre_archivo
            word_file_obj.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            upload_file_to_drive(service, word_file_obj, subapartado_guion_folder_id)
            if show_toast: st.toast(f"Borrador (GPT) para '{titulo}' generado y guardado.")
            return True
        except Exception as e: st.error(f"Error al generar con GPT para '{titulo}': {e}"); return False

    def ejecutar_regeneracion(titulo, file_id_borrador):
        st.warning(f"La función de re-generación para '{titulo}' aún no está implementada.")

    def ejecutar_borrado(titulo, folder_id_to_delete):
        st.warning(f"La función de borrado para '{titulo}' aún no está implementada.")

    # --- UI Y LÓGICA DE LA PÁGINA ---
    with st.spinner("Sincronizando con Google Drive..."):
        guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
        carpetas_existentes_response = get_files_in_project(service, guiones_folder_id)
        carpetas_de_guiones_existentes = {f['name']: f['id'] for f in carpetas_existentes_response if f['mimeType'] == 'application/vnd.google-apps.folder'}
        nombres_carpetas_existentes = set(carpetas_de_guiones_existentes.keys())

    st.subheader("Generación de Borradores en Lote")
    pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) not in nombres_carpetas_existentes]
    
    def toggle_all_checkboxes():
        new_state = st.session_state.select_all_checkbox
        for key in pending_keys: st.session_state[f"cb_{key}"] = new_state

    with st.container(border=True):
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_checkbox", on_change=toggle_all_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"cb_{key}")]
            num_selected = len(selected_keys)
            if st.button(f"🚀 Generar {num_selected} borradores seleccionados", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                progress_bar = st.progress(0, text="Iniciando generación en lote...")
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                for i, matiz_a_generar in enumerate(items_to_generate):
                    titulo = matiz_a_generar.get('subapartado')
                    progress_text = f"Generando ({i+1}/{num_selected}): {titulo}"
                    progress_bar.progress((i + 1) / num_selected, text=progress_text)
                    if USE_GPT_MODEL:
                        ejecutar_generacion_con_gpt(titulo, matiz_a_generar, show_toast=False)
                    else:
                        # <<< CORRECCIÓN AQUÍ 1 de 2 >>>
                        # Añadimos 'model' como primer argumento
                        ejecutar_generacion_con_gemini(model, titulo, matiz_a_generar, show_toast=False)
                progress_bar.progress(1.0, text="¡Generación en lote completada!")
                st.success(f"{num_selected} borradores generados.")
                st.balloons(); time.sleep(2); st.rerun()

    st.markdown("---")
    st.subheader("Gestión de Guiones de Subapartados")
    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get('subapartado')
        if not subapartado_titulo: continue
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
        
        if nombre_limpio in nombres_carpetas_existentes:
            estado = "📄 Generado"
            subapartado_folder_id = carpetas_de_guiones_existentes[nombre_limpio]
            files_in_subfolder = get_files_in_project(service, subapartado_folder_id)
            file_info = next((f for f in files_in_subfolder if f['name'].endswith('.docx')), None)
        else:
            estado = "⚪ No Generado"
            file_info, subapartado_folder_id = None, None
            
        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if estado == "⚪ No Generado": st.checkbox(f"**{subapartado_titulo}**", key=f"cb_{subapartado_titulo}")
                else: st.write(f"**{subapartado_titulo}**")
                st.caption(f"Estado: {estado}")
                if estado == "⚪ No Generado":
                    st.file_uploader("Aportar documentación de apoyo", type=['pdf', 'docx', 'txt'], key=f"upload_{subapartado_titulo}", accept_multiple_files=True, label_visibility="collapsed")
            with col2:
                if estado == "📄 Generado" and file_info:
                    link = f"https://docs.google.com/document/d/{file_info['id']}/edit"
                    st.link_button("Revisar en Drive", link, use_container_width=True)
                    if st.button("Re-Generar con Feedback", key=f"regen_{i}", type="primary", use_container_width=True):
                        ejecutar_regeneracion(subapartado_titulo, file_info['id'])
                    if st.button("🗑️ Borrar", key=f"del_{i}", use_container_width=True):
                        ejecutar_borrado(subapartado_titulo, subapartado_folder_id)
                else:
                    if st.button("Generar Borrador", key=f"gen_{i}", use_container_width=True):
                        with st.spinner(f"Generando borrador para '{subapartado_titulo}'..."):
                            if USE_GPT_MODEL:
                                if ejecutar_generacion_con_gpt(subapartado_titulo, matiz):
                                    st.rerun()
                            else:
                                # <<< CORRECCIÓN AQUÍ 2 de 2 >>>
                                # Añadimos 'model' como primer argumento
                                if ejecutar_generacion_con_gemini(model, subapartado_titulo, matiz):
                                    st.rerun()
                                    
    # --- NAVEGACIÓN ---
    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Revisión de Índice (F1)", on_click=go_to_phase1_results, use_container_width=True)
    with col_nav2:
        st.button("Ir a Plan de Prompts (F3) →", on_click=go_to_phase3, use_container_width=True)

def phase_3_page(model):
    """Página interactiva para generar, borrar, descargar y unificar planes de prompts."""
    from app import go_to_phase1, go_to_phase2, go_to_phase4
    from prompts import PROMPT_DESARROLLO

    st.markdown("<h3>FASE 3: Centro de Mando de Prompts</h3>", unsafe_allow_html=True)
    st.markdown("Genera planes de prompts de forma individual o selecciónalos para procesarlos en lote.")
    st.markdown("---")

    # --- SETUP INICIAL Y CARGA DE ÍNDICE ---
    service = st.session_state.drive_service
    if not service:
        st.error("No se pudo restablecer la conexión con Google Drive. Por favor, refresca la página.")
        st.stop()
    project_folder_id = st.session_state.selected_project['id']
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)

    if 'generated_structure' not in st.session_state:
        st.info("Sincronizando índice desde Google Drive...")
        saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
        if saved_index_id:
            index_content_bytes = download_file_from_drive(service, saved_index_id)
            st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
            st.rerun()
        else:
            st.warning("No se ha encontrado un índice. Vuelve a Fase 1 para generarlo.")
            if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
            return

    # --- CONSTRUCCIÓN DE LISTA ROBUSTA ---
    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict) and 'subapartado' in item}
    if not estructura: st.error("La estructura JSON no contiene la clave 'estructura_memoria'."); return
    subapartados_a_mostrar = []
    for seccion in estructura:
        apartado_principal = seccion.get('apartado', 'Sin Título')
        for subapartado_titulo in seccion.get('subapartados', []):
            matiz_existente = matices_dict.get(subapartado_titulo)
            if matiz_existente: subapartados_a_mostrar.append(matiz_existente)
            else: subapartados_a_mostrar.append({"apartado": apartado_principal, "subapartado": subapartado_titulo, "indicaciones": "No se encontraron indicaciones detalladas."})
    if not subapartados_a_mostrar: st.warning("El índice no contiene subapartados."); return

    # --- FUNCIONES DE ACCIÓN INTERNAS ---
    def handle_individual_generation(matiz_info, callback_model, show_toast=True):
        apartado_titulo = matiz_info.get("apartado", "N/A")
        subapartado_titulo = matiz_info.get("subapartado", "N/A")
        json_limpio_str = ""
        try:
            guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
            subapartado_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_main_folder_id)
            
            contexto_adicional_str = ""
            files_in_subfolder = get_files_in_project(service, subapartado_folder_id)
            for file_info in files_in_subfolder:
                file_bytes = download_file_from_drive(service, file_info['id'])
                if file_info['name'].endswith('.docx'):
                    doc = docx.Document(io.BytesIO(file_bytes.getvalue()))
                    texto_doc = "\n".join([p.text for p in doc.paragraphs])
                    contexto_adicional_str += f"\n--- CONTENIDO DEL GUION ({file_info['name']}) ---\n{texto_doc}\n"
                elif file_info['name'].endswith('.pdf'):
                    reader = PdfReader(io.BytesIO(file_bytes.getvalue()))
                    texto_pdf = "".join(page.extract_text() for page in reader.pages)
                    contexto_adicional_str += f"\n--- CONTENIDO DEL PDF DE APOYO ({file_info['name']}) ---\n{texto_pdf}\n"
            
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            pliegos_files_info = get_files_in_project(service, pliegos_folder_id)
            pliegos_content_for_ia = [{"mime_type": f['mimeType'], "data": download_file_from_drive(service, f['id']).getvalue()} for f in pliegos_files_info]
            
            prompt_final = PROMPT_DESARROLLO.format(apartado_titulo=apartado_titulo, subapartado_titulo=subapartado_titulo, indicaciones=matiz_info.get("indicaciones", ""))
            
            contenido_ia = [prompt_final] + pliegos_content_for_ia
            if contexto_adicional_str: 
                contenido_ia.append("--- CONTEXTO ADICIONAL DE GUIONES Y DOCUMENTACIÓN DE APOYO ---\n" + contexto_adicional_str)
            
            generation_config = {"response_mime_type": "application/json"}
            response = callback_model.generate_content(contenido_ia, generation_config=generation_config)
            
            json_limpio_str = limpiar_respuesta_json(response.text)
            if json_limpio_str:
                plan_parcial_obj = json.loads(json_limpio_str)
                json_bytes = json.dumps(plan_parcial_obj, indent=2, ensure_ascii=False).encode('utf-8')
                
                mock_file_obj = io.BytesIO(json_bytes)
                mock_file_obj.name = "prompts_individual.json"
                mock_file_obj.type = "application/json"
                
                old_plan_id = find_file_by_name(service, "prompts_individual.json", subapartado_folder_id)
                if old_plan_id: delete_file_from_drive(service, old_plan_id)
                
                upload_file_to_drive(service, mock_file_obj, subapartado_folder_id)
                if show_toast: st.toast(f"Plan para '{subapartado_titulo}' guardado.")
                return True
        except json.JSONDecodeError as json_err:
             st.error(f"Error Crítico: La IA devolvió un JSON inválido para '{subapartado_titulo}' que no se pudo reparar. Detalles: {json_err}")
             st.code(json_limpio_str)
             return False
        except Exception as e:
            st.error(f"Error generando prompts para '{subapartado_titulo}': {e}")
            return False

    def handle_individual_deletion(titulo, plan_id_to_delete):
        """Elimina un archivo de plan individual y refresca la página."""
        with st.spinner(f"Eliminando el plan para '{titulo}'..."):
            if delete_file_from_drive(service, plan_id_to_delete):
                st.toast(f"Plan para '{titulo}' eliminado con éxito.")
                st.rerun()

    def handle_conjunto_generation():
        """Unifica todos los planes individuales en un único archivo maestro."""
        with st.spinner("Unificando todos los planes de prompts..."):
            try:
                guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
                carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
                
                plan_conjunto_final = {"plan_de_prompts": []}
                
                for nombre_carpeta, folder_id in carpetas_de_guiones.items():
                    plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
                    if plan_id:
                        json_bytes = download_file_from_drive(service, plan_id).getvalue()
                        plan_individual_obj = json.loads(json_bytes.decode('utf-8'))
                        prompts_de_este_plan = plan_individual_obj.get("plan_de_prompts", [])
                        plan_conjunto_final["plan_de_prompts"].extend(prompts_de_este_plan)

                if not plan_conjunto_final["plan_de_prompts"]:
                    st.warning("No se encontraron planes individuales para unificar. Genera al menos uno.")
                    return

                nombre_archivo_final = "plan_de_prompts_conjunto.json"
                json_bytes_finales = json.dumps(plan_conjunto_final, indent=2, ensure_ascii=False).encode('utf-8')
                
                mock_file_obj = io.BytesIO(json_bytes_finales)
                mock_file_obj.name = nombre_archivo_final
                mock_file_obj.type = "application/json"
                
                old_conjunto_id = find_file_by_name(service, nombre_archivo_final, docs_app_folder_id)
                if old_conjunto_id: delete_file_from_drive(service, old_conjunto_id)
                
                upload_file_to_drive(service, mock_file_obj, docs_app_folder_id)
                st.success(f"¡Plan conjunto generado! Se unificaron {len(plan_conjunto_final['plan_de_prompts'])} prompts.")
                st.balloons()
            except Exception as e:
                st.error(f"Ocurrió un error durante la unificación: {e}")

    # OPTIMIZACIÓN: OBTENER ESTADO DE PLANES UNA SOLA VEZ
    with st.spinner("Verificando estado de los planes de prompts..."):
        guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
        carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
        planes_individuales_existentes = {}
        for nombre_carpeta, folder_id in carpetas_de_guiones.items():
            plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
            if plan_id: planes_individuales_existentes[nombre_carpeta] = plan_id

    # SECCIÓN SUPERIOR PARA ACCIONES EN LOTE
    st.subheader("Generación de Planes de Prompts en Lote")
    pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) in carpetas_de_guiones and re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) not in planes_individuales_existentes]
    
    def toggle_all_prompt_checkboxes():
        new_state = st.session_state.select_all_prompts_checkbox
        for key in pending_keys: st.session_state[f"pcb_{key}"] = new_state
        
    with st.container(border=True):
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_prompts_checkbox", on_change=toggle_all_prompt_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"pcb_{key}")]
            num_selected = len(selected_keys)
            if st.button(f"🚀 Generar {num_selected} planes seleccionados", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                progress_bar = st.progress(0, text="Iniciando generación en lote...")
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                for i, matiz_a_generar in enumerate(items_to_generate):
                    titulo = matiz_a_generar.get('subapartado')
                    progress_text = f"Generando plan ({i+1}/{num_selected}): {titulo}"
                    progress_bar.progress((i + 1) / num_selected, text=progress_text)
                    if not handle_individual_generation(matiz_a_generar, model, show_toast=False):
                        break # Detener en caso de error
                else: # Solo se ejecuta si el bucle termina sin break
                    progress_bar.progress(1.0, text="¡Generación en lote completada!")
                    st.success(f"{num_selected} planes generados.")
                    st.balloons()
                    st.rerun()

    st.markdown("---")
    st.subheader("Gestión de Planes de Prompts")

    # INTERFAZ DE GESTIÓN (HÍBRIDA Y OPTIMIZADA)
    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get("subapartado");
        if not subapartado_titulo: continue
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
        guion_generado = nombre_limpio in carpetas_de_guiones
        plan_individual_id = planes_individuales_existentes.get(nombre_limpio)
        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if not plan_individual_id and guion_generado:
                    st.checkbox(f"**{subapartado_titulo}**", key=f"pcb_{subapartado_titulo}")
                else: st.write(f"**{subapartado_titulo}**")
                
                if not guion_generado: st.warning("⚠️ Guion no generado en Fase 2. No se puede crear un plan.")
                elif plan_individual_id:
                    st.success("✔️ Plan generado")
                    with st.expander("Ver / Descargar Plan Individual"):
                        json_bytes = download_file_from_drive(service, plan_individual_id).getvalue()
                        st.json(json_bytes.decode('utf-8'))
                        st.download_button("Descargar JSON", data=json_bytes, file_name=f"prompts_{nombre_limpio}.json", mime="application/json", key=f"dl_{i}")
                else: st.info("⚪ Pendiente de generar plan de prompts")
            with col2:
                if not plan_individual_id:
                    st.button("Generar Plan de Prompts", key=f"gen_ind_{i}", on_click=handle_individual_generation, args=(matiz, model, True), use_container_width=True, type="primary", disabled=not guion_generado)
                else:
                    st.button("Re-generar Plan", key=f"gen_regen_{i}", on_click=handle_individual_generation, args=(matiz, model, True), use_container_width=True, type="secondary")
                    st.button("🗑️ Borrar Plan", key=f"del_plan_{i}", on_click=handle_individual_deletion, args=(subapartado_titulo, plan_individual_id), use_container_width=True)

    # BOTONES DE NAVEGACIÓN Y ACCIÓN FINAL
    st.markdown("---")
    st.button("🚀 Unificar y Guardar Plan de Prompts Conjunto", on_click=handle_conjunto_generation, use_container_width=True, type="primary", help="Unifica todos los planes individuales generados en un único archivo maestro.")
    col_nav3_1, col_nav3_2 = st.columns(2)
    with col_nav3_1:
        st.button("← Volver al Centro de Mando (F2)", on_click=go_to_phase2, use_container_width=True)
    with col_nav3_2:
        st.button("Ir a Redacción Final (F4) →", on_click=go_to_phase4, use_container_width=True)

def phase_4_page(model):
    """Página para ejecutar el plan de prompts y generar el cuerpo principal del documento Word."""
    from app import go_to_phase3, go_to_phase5
    st.markdown("<h3>FASE 4: Redacción del Cuerpo del Documento</h3>", unsafe_allow_html=True)
    st.markdown("Ejecuta el plan de prompts para generar el contenido completo de la memoria técnica. Este será el cuerpo principal del documento final.")
    st.markdown("---")

    # --- Setup inicial ---
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
    plan_conjunto_id = find_file_by_name(service, "plan_de_prompts_conjunto.json", docs_app_folder_id)

    if not plan_conjunto_id:
        st.warning("No se ha encontrado un 'plan_de_prompts_conjunto.json'. Vuelve a la Fase 3 para generarlo.")
        if st.button("← Ir a Fase 3"): go_to_phase3(); st.rerun()
        return

    try:
        json_bytes = download_file_from_drive(service, plan_conjunto_id).getvalue()
        plan_de_accion = json.loads(json_bytes.decode('utf-8'))
        lista_de_prompts = plan_de_accion.get("plan_de_prompts", [])
        lista_de_prompts.sort(key=lambda x: x.get('prompt_id', ''))
        st.success(f"✔️ Plan de acción cargado. Se ejecutarán {len(lista_de_prompts)} prompts para crear el cuerpo del documento.")
    except Exception as e:
        st.error(f"Error al cargar o procesar el plan de acción: {e}"); return

    # --- Lógica del botón de generación ---
    button_text = "🔁 Volver a Generar Cuerpo del Documento" if st.session_state.get("generated_doc_buffer") else "🚀 Iniciar Redacción y Generar Cuerpo"
    if st.button(button_text, type="primary", use_container_width=True):
        if not lista_de_prompts:
            st.warning("El plan de acción está vacío."); return

        generation_successful = False
        documento = docx.Document()
        try:
            with st.spinner("Iniciando redacción... Esto puede tardar varios minutos."):
                chat_redaccion = model.start_chat()
                progress_bar = st.progress(0, text="Configurando sesión de chat...")
                ultimo_apartado_escrito = None
                ultimo_subapartado_escrito = None
                for i, tarea in enumerate(lista_de_prompts):
                    progress_text = f"Procesando Tarea {i+1}/{len(lista_de_prompts)}: {tarea.get('subapartado_referencia', 'N/A')}"
                    progress_bar.progress((i + 1) / len(lista_de_prompts), text=progress_text)
                    
                    # --- LÓGICA DE ENCABEZADOS MEJORADA ---
                    apartado_actual = tarea.get("apartado_referencia")
                    subapartado_actual = tarea.get("subapartado_referencia")
                    if apartado_actual and apartado_actual != ultimo_apartado_escrito:
                        if ultimo_apartado_escrito is not None: documento.add_page_break()
                        documento.add_heading(apartado_actual, level=1)
                        ultimo_apartado_escrito = apartado_actual
                        ultimo_subapartado_escrito = None # Resetea el subapartado al cambiar de apartado
                    
                    if subapartado_actual and subapartado_actual != ultimo_subapartado_escrito:
                        documento.add_heading(subapartado_actual, level=2)
                        ultimo_subapartado_escrito = subapartado_actual

                    respuesta_ia_bruta = ""
                    prompt_actual = tarea.get("prompt_para_asistente")
                    
                    if prompt_actual:
                        response = chat_redaccion.send_message(prompt_actual)
                        respuesta_ia_bruta = response.text
                        time.sleep(1) 

                    # ========== INICIO DE LA CORRECCIÓN CLAVE ==========
                    # Detectamos HTML si el prompt lo indica O si la respuesta empieza con código HTML
                    es_html = (
                        "HTML" in tarea.get("prompt_id", "").upper() or 
                        "VISUAL" in tarea.get("prompt_id", "").upper() or
                        respuesta_ia_bruta.strip().startswith(('<!DOCTYPE html>', '<div', '<table'))
                    )

                    if es_html:
                        # Es un elemento visual: convertir a imagen
                        html_puro = limpiar_respuesta_final(respuesta_ia_bruta) # Limpiamos por si la IA añade texto extra
                        image_file = html_a_imagen(wrap_html_fragment(html_puro), f"temp_img_{i}.png")
                        if image_file and os.path.exists(image_file):
                            documento.add_picture(image_file, width=docx.shared.Inches(6.5))
                            os.remove(image_file)
                        else:
                            documento.add_paragraph("[ERROR AL GENERAR IMAGEN DESDE HTML]")
                    else:
                        # Es texto narrativo: limpiar y añadir como Markdown
                        texto_limpio = limpiar_respuesta_final(respuesta_ia_bruta)
                        texto_corregido = corregir_numeracion_markdown(texto_limpio)
                        if texto_corregido:
                            agregar_markdown_a_word(documento, texto_corregido)
                    # ========== FIN DE LA CORRECCIÓN CLAVE ==========

                generation_successful = True
        except Exception as e:
            st.error(f"Ocurrió un error crítico durante la generación del cuerpo del documento: {e}")
        
        if generation_successful:
            project_name = st.session_state.selected_project['name']
            safe_project_name = re.sub(r'[\\/*?:"<>|]', "", project_name).replace(' ', '_')
            nombre_archivo_final = f"Cuerpo_Memoria_Tecnica_{safe_project_name}.docx"
            
            doc_io = io.BytesIO()
            documento.save(doc_io)
            doc_io.seek(0)
            
            st.session_state.generated_doc_buffer = doc_io
            st.session_state.generated_doc_filename = nombre_archivo_final
            
            st.success("¡Cuerpo del documento generado con éxito!")
            st.rerun()

    # --- SECCIÓN DE DESCARGA Y NAVEGACIÓN ---
    if st.session_state.get("generated_doc_buffer"):
        st.info("El cuerpo del documento está listo. Ahora puedes descargarlo o pasar a la fase final de ensamblaje.")
        st.download_button(
            label="📄 Descargar Cuerpo del Documento (.docx)",
            data=st.session_state.generated_doc_buffer,
            file_name=st.session_state.generated_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Fase 3 (Plan de Prompts)", on_click=go_to_phase3, use_container_width=True)
    with col_nav2:
        st.button("Ir a Ensamblaje Final (F5) →", on_click=go_to_phase5, use_container_width=True, type="primary", 
                  disabled=not st.session_state.get("generated_doc_buffer"))
        
def phase_5_page(model):
    """
    Fase final y segura que ensambla el documento definitivo:
    1. Crea un Índice automático.
    2. Genera una Introducción estratégica basada en el cuerpo del documento.
    3. Inserta estos dos elementos al principio del borrador intacto de la Fase 4.
    """
    from app import go_to_phase4, go_to_phase1, back_to_project_selection_and_cleanup
    from prompts import PROMPT_GENERAR_INTRODUCCION

    st.markdown("<h3>FASE 5: Ensamblaje del Documento Final</h3>", unsafe_allow_html=True)
    st.markdown("Este es el último paso. El asistente tomará el documento generado en la fase anterior y le añadirá un índice y una introducción profesional para crear la versión definitiva.")
    st.markdown("---")

    # Comprobaciones iniciales
    if not st.session_state.get("generated_doc_buffer"):
        st.warning("No se ha encontrado un documento de la Fase 4 para trabajar. Por favor, completa la fase anterior.")
        if st.button("← Ir a Fase 4"): go_to_phase4(); st.rerun()
        return
    if not st.session_state.get("generated_structure"):
        st.warning("No se ha encontrado la estructura del proyecto. Vuelve a la Fase 1.")
        if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
        return

    if st.button("🚀 Ensamblar Documento Final con Índice e Introducción", type="primary", use_container_width=True):
        try:
            with st.spinner("Ensamblando la versión definitiva..."):
                # --- PREPARACIÓN ---
                buffer_fase4 = st.session_state.generated_doc_buffer
                buffer_fase4.seek(0)
                documento_fase4 = docx.Document(buffer_fase4)
                
                texto_completo_original = "\n".join([p.text for p in documento_fase4.paragraphs if p.text.strip()])

                # --- GENERACIÓN DE PIEZAS NUEVAS ---
                st.toast("Generando introducción estratégica...")
                response_intro = model.generate_content([PROMPT_GENERAR_INTRODUCCION, texto_completo_original])
                introduccion_markdown = limpiar_respuesta_final(response_intro.text)

                # --- ENSAMBLAJE FINAL ---
                st.toast("Creando documento final...")
                documento_final = docx.Document()
                
                # 1. Añadir el ÍNDICE
                estructura_memoria = st.session_state.generated_structure.get('estructura_memoria', [])
                generar_indice_word(documento_final, estructura_memoria)
                documento_final.add_page_break()
                
                # 2. Añadir la INTRODUCCIÓN
                documento_final.add_heading("Introducción", level=1)
                agregar_markdown_a_word(documento_final, corregir_numeracion_markdown(introduccion_markdown))
                documento_final.add_page_break()
                
                # 3. Añadir el CUERPO del documento de la Fase 4
                for element in documento_fase4.element.body:
                    documento_final.element.body.append(element)

                # --- GUARDADO Y DESCARGA ---
                doc_io_final = io.BytesIO()
                documento_final.save(doc_io_final)
                doc_io_final.seek(0)

                st.session_state.refined_doc_buffer = doc_io_final
                original_filename = st.session_state.generated_doc_filename
                # Creamos un nombre de archivo más adecuado
                st.session_state.refined_doc_filename = original_filename.replace("Cuerpo_", "Version_Final_")
                
                st.success("¡Documento final ensamblado con éxito!")
                st.rerun()

        except Exception as e:
            st.error(f"Ocurrió un error crítico durante el ensamblaje final: {e}")

    # Lógica de descarga
    if st.session_state.get("refined_doc_buffer"):
        st.balloons()
        st.success("¡Tu memoria técnica definitiva está lista!")
        st.download_button(
            label="🏆 Descargar Versión Definitiva (.docx)",
            data=st.session_state.refined_doc_buffer,
            file_name=st.session_state.refined_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    # Navegación
    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Fase 4", on_click=go_to_phase4, use_container_width=True)
    with col_nav2:

        st.button("↩️ Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup, use_container_width=True)

















