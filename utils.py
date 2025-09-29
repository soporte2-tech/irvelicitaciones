import re
import os
import io
import docx
import streamlit as st
import imgkit
from pypdf import PdfReader

# --- FUNCIONES AUXILIARES DE BACKEND ---
def limpiar_respuesta_json(texto_sucio):
    if not isinstance(texto_sucio, str): return ""
    match_bloque = re.search(r'```(?:json)?\s*(\{.*\})\s*```', texto_sucio, re.DOTALL)
    if match_bloque: return match_bloque.group(1).strip()
    match_objeto = re.search(r'\{.*\}', texto_sucio, re.DOTALL)
    if match_objeto: return match_objeto.group(0).strip()
    return ""

def agregar_markdown_a_word(documento, texto_markdown):
    patron_encabezado = re.compile(r'^(#+)\s+(.*)')
    patron_lista_numerada = re.compile(r'^\s*\d+\.\s+')
    patron_lista_viñeta = re.compile(r'^\s*[\*\-]\s+')
    def procesar_linea_con_negritas(parrafo, texto):
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'): parrafo.add_run(parte[2:-2]).bold = True
            elif parte: parrafo.add_run(parte)
    for linea in texto_markdown.split('\n'):
        linea_limpia = linea.strip()
        if not linea_limpia: continue
        match_encabezado = patron_encabezado.match(linea_limpia)
        if match_encabezado:
            documento.add_heading(match_encabezado.group(2).strip(), level=min(len(match_encabezado.group(1)), 4))
            continue
        if patron_lista_numerada.match(linea_limpia):
            p = documento.add_paragraph(style='List Number')
            procesar_linea_con_negritas(p, patron_lista_numerada.sub('', linea_limpia))
        elif patron_lista_viñeta.match(linea_limpia):
            p = documento.add_paragraph(style='List Bullet')
            procesar_linea_con_negritas(p, patron_lista_viñeta.sub('', linea_limpia))
        else:
            p = documento.add_paragraph()
            procesar_linea_con_negritas(p, linea_limpia)

def mostrar_indice_desplegable(estructura_memoria):
    if not estructura_memoria:
        st.warning("No se encontró una estructura de memoria para mostrar.")
        return
    st.subheader("Índice Propuesto")
    for seccion in estructura_memoria:
        apartado_titulo = seccion.get("apartado", "Apartado sin título")
        subapartados = seccion.get("subapartados", [])
        with st.expander(f"**{apartado_titulo}**"):
            if subapartados:
                for sub in subapartados: st.markdown(f"- {sub}")
            else: st.markdown("_Este apartado no tiene subapartados definidos._")

def sanitize_json_string(json_str):
    """
    Elimina caracteres de control inválidos de un string antes de parsearlo como JSON.
    Estos caracteres a veces son introducidos por el LLM al procesar PDFs/DOCX.
    """
    # Expresión regular para encontrar caracteres de control ASCII (0-31),
    # excepto los que son válidos en JSON strings si están escapados (tab, newline, etc.).
    # Esta regex busca los que causan errores de parseo.
    control_chars_regex = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')
    
    # Reemplazamos los caracteres problemáticos por una cadena vacía.
    sanitized_str = control_chars_regex.sub('', json_str)
    return sanitized_str

def wrap_html_fragment(html_fragment):
    """
    Toma un fragmento de HTML (ej: un <table> o <div>) y lo envuelve en una
    estructura HTML completa con los estilos CSS necesarios para renderizarlo.
    """
    # Si ya es un documento completo, lo devuelve tal cual.
    if html_fragment.strip().startswith('<!DOCTYPE html>'):
        return html_fragment

    # Estilos CSS extraídos de tu PROMPT_DESARROLLO.
    # Son necesarios para que las tablas y cards se vean bien.
    css_styles = """
        @import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');
        body {
            font-family: 'Urbanist', sans-serif;
            background-color: #f0f2f5;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
            width: 800px;
            box-sizing: border-box;
        }
        .card {
            background-color: white;
            border-radius: 10px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            padding: 25px;
            width: 100%;
            max-width: 750px;
            border-top: 5px solid #0046C6;
        }
        h2 {
            color: #0046C6;
            text-align: center;
            margin-top: 0;
            font-size: 24px;
            font-weight: 700;
        }
        ul { list-style-type: none; padding: 0; }
        li {
            display: flex;
            align-items: center;
            margin-bottom: 15px;
            font-size: 16px;
            color: #333;
        }
        li::before {
            content: '✔';
            color: #32CFAA;
            font-size: 20px;
            font-weight: bold;
            margin-right: 15px;
        }
        /* Estilos adicionales para tablas */
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            font-size: 15px;
        }
        th, td {
            padding: 12px 15px;
            border: 1px solid #ddd;
            text-align: left;
        }
        th {
            background-color: #f5f5f5;
            font-weight: 600;
            color: #333;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
    """
    
    # Plantilla HTML completa
    full_html_template = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Visual Element</title>
        <style>{css_styles}</style>
    </head>
    <body>
        {html_fragment}
    </body>
    </html>
    """
    return full_html_template

def html_a_imagen(html_string, output_filename="temp_image.png"):
    """
    Convierte una cadena de HTML en una imagen PNG, encontrando automáticamente
    el ejecutable wkhtmltoimage en el entorno de Streamlit Cloud.
    """
    try:
        # En Streamlit Cloud, el ejecutable se instala en una ruta accesible.
        # 'which' es un comando de Linux para encontrar la ruta de un programa.
        path_wkhtmltoimage = os.popen('which wkhtmltoimage').read().strip()

        if not path_wkhtmltoimage:
            st.error("❌ El ejecutable 'wkhtmltoimage' no se encontró. Asegúrate de que 'wkhtmltopdf' está en tu packages.txt y que la app ha sido reiniciada.")
            return None

        # Crea una configuración para imgkit apuntando al ejecutable encontrado.
        config = imgkit.config(wkhtmltoimage=path_wkhtmltoimage)
        
        # Opciones para mejorar la calidad y el tamaño de la imagen
        options = {
            'format': 'png',
            'encoding': "UTF-8",
            'width': '800',  # Un ancho fijo para consistencia
            'quiet': ''      # Suprime la salida de la consola
        }

        # Genera la imagen desde la cadena de HTML
        imgkit.from_string(html_string, output_filename, config=config, options=options)
        
        if os.path.exists(output_filename):
            return output_filename
        else:
            st.warning(f"imgkit ejecutado pero el archivo '{output_filename}' no fue creado.")
            return None

    except Exception as e:
        st.error(f"Ocurrió un error al convertir HTML a imagen: {e}")
        st.code(f"Path de wkhtmltoimage intentado: {os.popen('which wkhtmltoimage').read().strip()}", language="bash")
        return None

# AÑADE ESTA NUEVA FUNCIÓN A TU SCRIPT
def limpiar_respuesta_narrativa(texto_ia):
    """
    Limpia la respuesta de la IA para la redacción final, eliminando:
    - Bloques de código JSON.
    - Frases introductorias comunes.
    - El propio título del subapartado si la IA lo repite.
    """
    if not isinstance(texto_ia, str):
        return ""

    # Eliminar bloques de código JSON completos
    texto_limpio = re.sub(r'```json\s*\{.*?\}\s*```', '', texto_ia, flags=re.DOTALL)
    
    # Eliminar frases introductorias comunes (puedes añadir más)
    frases_a_eliminar = [
        r'^\s*Aquí tienes el contenido para el subapartado.*?:',
        r'^\s*Claro, aquí está la redacción para.*?:',
        r'^\s*A continuación se presenta el contenido detallado:',
        r'^\s*##\s*.*?$' # Elimina cualquier título Markdown que la IA pueda añadir
    ]
    for patron in frases_a_eliminar:
        texto_limpio = re.sub(patron, '', texto_limpio, flags=re.IGNORECASE | re.MULTILINE).strip()

    return texto_limpio

def corregir_numeracion_markdown(texto_markdown):
    """
    Recorre un texto en Markdown y corrige las listas numeradas para que
    siempre empiecen en 1 y sean consecutivas.
    """
    lineas_corregidas = []
    contador_lista = 0
    en_lista_numerada = False

    for linea in texto_markdown.split('\n'):
        # Usamos una regex para detectar si la línea empieza como un item de lista numerada
        match = re.match(r'^\s*\d+\.\s+', linea)
        if match:
            if not en_lista_numerada:
                # Si es el primer item de una nueva lista, reiniciamos el contador
                en_lista_numerada = True
                contador_lista = 1
            else:
                # Si ya estábamos en una lista, incrementamos
                contador_lista += 1

            # Reemplazamos el número original por el correcto
            texto_del_item = linea[match.end():]
            lineas_corregidas.append(f"{contador_lista}. {texto_del_item}")
        else:
            # Si la línea no es un item de lista, se resetea el estado
            en_lista_numerada = False
            contador_lista = 0
            lineas_corregidas.append(linea)

    return '\n'.join(lineas_corregidas)

# AÑADE ESTA NUEVA FUNCIÓN A TUS FUNCIONES AUXILIARES

def generar_indice_word(documento, estructura_memoria):
    """
    Añade un índice (Tabla de Contenidos) al principio de un documento de Word
    basado en la estructura de la memoria técnica.
    """
    documento.add_heading("Índice", level=1)
    
    if not estructura_memoria:
        documento.add_paragraph("No se encontró una estructura para generar el índice.")
        return

    for seccion in estructura_memoria:
        apartado_titulo = seccion.get("apartado", "Apartado sin título")
        subapartados = seccion.get("subapartados", [])
        
        # Añade el apartado principal
        p = documento.add_paragraph()
        p.add_run(apartado_titulo).bold = True
        
        # Añade los subapartados con sangría
        if subapartados:
            for sub in subapartados:
                # Usamos un estilo de párrafo con sangría si existe, o añadimos espacios
                p_sub = documento.add_paragraph(f"    {sub}")

    st.toast("Índice generado en el documento.")

def limpiar_respuesta_final(texto_ia):
    """
    Limpia de forma agresiva la respuesta de la IA, eliminando todo
    el "meta-texto", explicaciones, y bloques de código mal formateados.
    """
    if not isinstance(texto_ia, str):
        return ""

    # <-- ¡NUEVA REGLA! Elimina comentarios específicos sobre la creación de diagramas/código.
    # Esto busca frases que empiezan con "Este código..." y terminan con "...visualizar el diagrama." y lo elimina todo.
    texto_limpio = re.sub(r'Este código crea.*?visualizar el diagrama\.', '', texto_ia, flags=re.DOTALL | re.IGNORECASE)
    
    # Eliminar explicaciones comunes sobre el código HTML que la IA añade al final
    texto_limpio = re.sub(r'El código HTML proporcionado genera.*?aún más:', '', texto_limpio, flags=re.DOTALL | re.IGNORECASE)
    
    # Eliminar cualquier bloque de código JSON que pueda haberse colado
    texto_limpio = re.sub(r'```json\s*\{.*?\}\s*```', '', texto_limpio, flags=re.DOTALL)

    # Eliminar los marcadores de bloque de código de texto plano o html
    texto_limpio = re.sub(r'```(html|mermaid)?', '', texto_limpio)
    
    # Eliminar frases introductorias o de cierre que a veces añade la IA
    frases_a_eliminar = [
        r'^\s*Aquí tienes el contenido.*?:',
        r'^\s*Claro, aquí está la redacción para.*?:',
        r'^\s*##\s*.*?$' # Elimina cualquier título Markdown que la IA pueda repetir
    ]
    for patron in frases_a_eliminar:
        texto_limpio = re.sub(patron, '', texto_limpio, flags=re.IGNORECASE | re.MULTILINE)

    return texto_limpio.strip()